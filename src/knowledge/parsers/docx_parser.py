"""
DOCX parser — extracts heading hierarchy and tables from .docx files.

Uses python-docx to map Heading 1-6 styles to section levels.
Adapted from Scriptiva's passage_extraction hierarchy patterns.
"""

import logging
from pathlib import Path
from typing import List, Optional

from .models import ParsedDocument, DocumentSection, TableData, ParseError

logger = logging.getLogger(__name__)

# Map Word heading styles to levels
_HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
    "Title": 1,
    "Subtitle": 2,
}


class DocxParser:
    """Parse DOCX documents extracting heading hierarchy and tables."""

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a .docx file into a ParsedDocument."""
        if not file_path.exists():
            raise ParseError(f"File not found: {file_path}")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ParseError("python-docx is required for DOCX parsing. Install with: pip install python-docx")

        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            raise ParseError(f"Failed to open DOCX: {exc}")

        sections = self._extract_sections(doc)
        tables = self._extract_tables(doc)

        # Determine title
        title = file_path.stem
        if doc.core_properties.title:
            title = doc.core_properties.title
        elif sections and sections[0].heading:
            title = sections[0].heading

        parsed = ParsedDocument(
            source_path=str(file_path),
            title=title,
            sections=sections,
            tables=tables,
            metadata={
                "parser": "python-docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )
        parsed.compute_hash()
        logger.info(
            "Parsed %s (%d sections, %d tables)",
            file_path.name,
            len(sections),
            len(tables),
        )
        return parsed

    def _extract_sections(self, doc: "DocxDocument") -> List[DocumentSection]:
        """Extract sections from paragraphs, mapping heading styles to levels."""
        sections: List[DocumentSection] = []
        current = DocumentSection(heading="", level=0, content="")

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            level = _HEADING_LEVELS.get(style_name, 0)

            if level > 0:
                # New heading — flush current section
                if current.content or current.heading:
                    sections.append(current)
                current = DocumentSection(heading=text, level=level, content="")
            else:
                # Body text — append to current section
                if current.content:
                    current.content += "\n" + text
                else:
                    current.content = text

        # Flush last section
        if current.content or current.heading:
            sections.append(current)

        # If no headings found, wrap everything in a single section
        if not sections:
            all_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if all_text:
                sections.append(DocumentSection(heading="", level=0, content=all_text))

        return sections

    def _extract_tables(self, doc: "DocxDocument") -> List[TableData]:
        """Extract tables from the document."""
        tables: List[TableData] = []

        for table in doc.tables:
            rows_data: List[List[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)

            if not rows_data:
                continue

            # First row as headers if it looks like a header
            headers = rows_data[0] if rows_data else []
            data_rows = rows_data[1:] if len(rows_data) > 1 else []

            tables.append(TableData(headers=headers, rows=data_rows))

        return tables
